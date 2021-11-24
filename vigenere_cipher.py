import docx
from docx.shared import RGBColor, Pt

#get text from fake message, turn lines to list items
fake_text = docx.Document('fakeMessage.docx')
fake_list = []
for paragraph in fake_text.paragraphs:
    fake_list.append(paragraph.text)

#get text from real message, turn lines to list items
real_text = docx.Document('realMessage.docx')
real_list = []
for paragraph in real_text.paragraphs:
    #remove blank lines
    if len(paragraph.text) != 0:
        real_list.append(paragraph.text)

#load a template
doc = docx.Document('template.docx')

#fill in template
doc.add_heading('Morland Holmes', 0)
subtitle = doc.add_heading('Global Consulting & Negotiations', 1)
subtitle.alignment = 1
doc.add_heading('', 1)
doc.add_paragraph('December 17 2015')
doc.add_paragraph('')

#Use python-docx's paragraph_format to set the lines in order
def set_spacing(paragraph):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

length_real = len(real_list)
count_real = 0

#combine the real and fake messages
for line in fake_list:
    if count_real < length_real and line == "":
        paragraph = doc.add_paragraph(real_list[count_real])
        paragraph_index = len(doc.paragraphs) - 1
        #change real message color
        run = doc.paragraphs[paragraph_index].runs[0]
        font = run.font
        font.color.rgb = RGBColor(255, 255, 255)
        count_real += 1
    else:
        paragraph = doc.add_paragraph(line)

    set_spacing(paragraph)

doc.save('ciphertext_message_letterhead.docx')

print("Done")
